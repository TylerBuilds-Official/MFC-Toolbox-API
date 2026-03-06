"""
Extraction utilities for preprocessing uploaded files into LLM-friendly text.

# TODO: Add tool-based drilldown support for specific file types.
#       - xlsx: allow model to request specific sheets, cell ranges, or filtered views
#       - docx: allow model to request specific sections, tables, or embedded images
#       Register a `read_uploaded_file` tool that accepts file_id + optional params
#       (sheet_name, page_range, section_heading, etc.) for on-demand deep inspection.
"""

import io
from pathlib import Path

from PIL import Image
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from openpyxl import load_workbook


# Max base64 size for LLM APIs (Anthropic limit is 5MB)
LLM_IMAGE_MAX_BYTES = 4 * 1024 * 1024  # 4MB to leave headroom


def compress_image_for_llm(file_path: str, max_bytes: int = LLM_IMAGE_MAX_BYTES) -> tuple[bytes, str]:
    """Compress an image to fit within LLM API limits.

    Returns (compressed_bytes, mime_type). Keeps original format if small enough,
    otherwise converts to JPEG and scales down progressively.
    """

    raw = Path(file_path).read_bytes()
    if len(raw) <= max_bytes:

        # Infer mime from extension
        ext = Path(file_path).suffix.lower()
        mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                    ".gif": "image/gif", ".webp": "image/webp"}

        return raw, mime_map.get(ext, "image/jpeg")

    img = Image.open(file_path)
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")

    # Try quality reduction first
    for quality in (85, 70, 50, 30):
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=quality, optimize=True)
        if buf.tell() <= max_bytes:

            return buf.getvalue(), "image/jpeg"

    # Scale down if quality alone isn't enough
    for scale in (0.75, 0.5, 0.35, 0.25):
        resized = img.resize(
            (int(img.width * scale), int(img.height * scale)),
            Image.LANCZOS
        )
        buf = io.BytesIO()
        resized.save(buf, format="JPEG", quality=50, optimize=True)
        if buf.tell() <= max_bytes:

            return buf.getvalue(), "image/jpeg"

    # Last resort — tiny
    resized = img.resize((800, int(800 * img.height / img.width)), Image.LANCZOS)
    buf = io.BytesIO()
    resized.save(buf, format="JPEG", quality=40, optimize=True)

    return buf.getvalue(), "image/jpeg"


def _runs_to_markdown(para) -> str:
    """Convert paragraph runs to markdown with bold/italic."""

    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue

        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"

        parts.append(text)

    return "".join(parts).strip()


def extract_docx_as_markdown(file_path: str) -> str:
    """Extract a .docx file to markdown-flavored text."""

    doc    = DocxDocument(file_path)
    parts: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split('}')[-1]

        if tag == 'p':
            para = next(
                (p for p in doc.paragraphs if p._element is element),
                None
            )
            if para is None:
                continue

            text  = _runs_to_markdown(para)
            style = (para.style.name or '').lower()

            if not text:
                continue

            if 'heading 1' in style:
                parts.append(f"# {text}")
            elif 'heading 2' in style:
                parts.append(f"## {text}")
            elif 'heading 3' in style:
                parts.append(f"### {text}")
            elif 'heading' in style:
                parts.append(f"#### {text}")
            elif 'list' in style or 'bullet' in style:
                parts.append(f"- {text}")
            else:
                parts.append(text)

        elif tag == 'tbl':
            table = next(
                (t for t in doc.tables if t._element is element),
                None
            )
            if table:
                parts.append(_docx_table_to_markdown(table))

    return "\n\n".join(parts)


def _docx_table_to_markdown(table: DocxTable) -> str:
    """Convert a docx table to a markdown table."""

    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows.append(cells)

    if not rows:

        return ""

    lines = []

    # Header row
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")

    # Data rows
    for row in rows[1:]:
        # Pad row if fewer cells than header
        padded = row + [""] * (len(rows[0]) - len(row))
        lines.append("| " + " | ".join(padded) + " |")

    return "\n".join(lines)


def extract_xlsx_as_text(file_path: str, max_rows: int = 500) -> str:
    """Extract an .xlsx file to structured text, one section per sheet.

    NOTE: This only extracts cell data. Charts, conditional formatting, and
    other visual elements are lost. Future improvement: use win32com or
    LibreOffice headless to export a PDF companion that preserves visuals,
    then send both the structured text AND the PDF to the LLM.
    """

    wb     = load_workbook(file_path, read_only=True, data_only=True)
    parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws   = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))

        if not rows:
            parts.append(f"## Sheet: {sheet_name}\n(empty)")
            continue

        # Find header row (first non-empty row)
        headers    = [str(c) if c is not None else "" for c in rows[0]]
        data_rows  = rows[1:]
        total_rows = len(data_rows)
        truncated  = total_rows > max_rows

        if truncated:
            # Show first and last rows with a gap indicator
            head = data_rows[:max_rows // 2]
            tail = data_rows[-(max_rows // 2):]
            data_rows = head + [None] + tail

        lines = [f"## Sheet: {sheet_name} ({total_rows} rows)"]

        # Build markdown table
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

        for row in data_rows:
            if row is None:
                lines.append(f"| ... ({total_rows - max_rows} rows omitted) |")
                continue

            cells = [str(c) if c is not None else "" for c in row]
            # Pad if fewer cells than headers
            padded = cells + [""] * (len(headers) - len(cells))
            lines.append("| " + " | ".join(padded[:len(headers)]) + " |")

        parts.append("\n".join(lines))

    wb.close()

    return "\n\n".join(parts)
